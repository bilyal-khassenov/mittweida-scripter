# Import main packages
import pathlib, os, re, sys, time, datetime, shutil, torch, pandas, whisper, getpass, traceback, json, tiktoken
import ffmpeg
from mutagen.oggopus import OggOpus
from docx.enum.text import WD_COLOR_INDEX
from docx import Document
from docx.shared import Pt
from whisper.utils import get_writer
from ollama import chat
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
######from docx.document import Document #Keep it for Intellisense!
###### from docx.document import Document # Keep it for Intellisense!

# Import helper functions
import mws_helpers


# Central paths definition
new_line_for_f_strings = '\n'
dir_temp_orig_files = mws_helpers.ProjectPaths().temp_orig_file_path
dir_format_conversion = mws_helpers.ProjectPaths().folder_for_format_conversion_path
dir_in_progress = mws_helpers.ProjectPaths().in_progress_folder_path
dir_processed = mws_helpers.ProjectPaths().processed_folder_path
dir_orig_files_temps = mws_helpers.ProjectPaths().temp_orig_file_path
path_to_perf_protocol = mws_helpers.ProjectPaths().performance_protocol_fullfilename
configs = mws_helpers.get_configs()
enc = tiktoken.get_encoding("cl100k_base")

# Input budget per LLM call: leave headroom in the context window for the system
# prompt and the generated summary, and absorb the tokenizer mismatch between
# tiktoken (used for counting) and the actual model tokenizer
llm_input_token_budget = configs['llm']['num_ctx'] - 8192

#create logger
logger = mws_helpers.create_logger(__name__)

def notify_admins(message):
    if configs['telegram']['use_telegram'] == True:
        mws_helpers.send_telegram_message(
            configs['telegram']['admin_chat_id'],
            message
        )

def diarize_file(file_path):
    logger.debug("Diarizing file...")
    from pyannote.audio import Pipeline

    pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1")

    # Check if CUDA is available and if it is, send pipeline to GPU
    if torch.cuda.is_available():
        pipeline.to(torch.device("cuda"))

    # Apply pretrained pipeline
    diarization = pipeline(str(file_path))

    speech_turns = []
    index = 0

    for turn, _, speaker in diarization.itertracks(yield_label=True):
        speech_turns.append({
            'index': index,
            'start': turn.start,
            'end': turn.end,
            'speaker': speaker
        })
        index = index + 1

    logger.debug("Diarization finished successfully")
    return speech_turns


def diarize_timestamped_words(conversation_turns, timestamped_words):
    logger.debug("Diarizing timestamped words...")
    for conversation_turn in conversation_turns:
        # First prepare list of all words marked as temporary acceptable or not for this speech turn
        for word in timestamped_words:
            # Word must end after the start of conversation turn
            word['start_acceptable'] = float(word['end']) > float(conversation_turn['start'])

            # Word must start before the end of conversation turn
            word['end_acceptable'] = float(word['start']) < float(conversation_turn['end'])

            # Both previous conditions must be met
            word['acceptance_result'] = (
                word['start_acceptable'] == True
                and word['end_acceptable'] == True
            )

        # Collect conversation turn from individual words that were accepted for it
        conversation_turn_text = ''

        for word in timestamped_words:
            if word['acceptance_result'] == True:
                conversation_turn_text = conversation_turn_text + word['word']

        # Save the collected conversation turn text to the processed conversation turn
        conversation_turn['text'] = conversation_turn_text.strip()

    return conversation_turns


def word_count(text):
    return len(text.split())


def format_timestamp(seconds):
    return time.strftime('%H:%M:%S', time.gmtime(seconds))


def format_segments_for_llm(segments):
    return [segment['text'].strip() for segment in segments if segment['text'].strip()]


def format_conversation_turns_for_llm(conversation_turns):
    return [
        f"[{format_timestamp(conversation_turn['start'])}] "
        f"{conversation_turn['speaker']}: {conversation_turn['text'].strip()}"
        for conversation_turn in conversation_turns
    ]


def group_texts_by_tokens(texts, max_tokens):
    groups = []
    current_group = []
    current_tokens = 0
    for text in texts:
        # Count the joining newline as well, so joined groups stay within the budget
        text_tokens = len(enc.encode(text + "\n"))
        if current_group and current_tokens + text_tokens > max_tokens:
            groups.append(current_group)
            current_group = []
            current_tokens = 0
        current_group.append(text)
        current_tokens += text_tokens
    if current_group:
        groups.append(current_group)
    return groups


def chunk_lines_by_tokens(lines, max_tokens):
    return ["\n".join(group) for group in group_texts_by_tokens(lines, max_tokens)]


def clean_llm_output(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = text.replace('**', '')
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\*\s+', '- ', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def summarize_chunk(text, prompt_hint, summary_language, min_words, max_words, prompt_template):
    logger.debug(f"Summarizing chunk with target length {min_words}-{max_words} words")

    prompt = prompt_template.format(
        min_words=min_words,
        max_words=max_words,
        summary_language=summary_language
    )

    system_prompt = prompt + prompt_hint

    options = {
        "num_ctx": configs['llm']['num_ctx'],
        "temperature": configs['llm']['temperature'],
        "num_predict": min(8192, max(768, int(max_words * 3)))
    }

    response = chat(
        model=configs['llm']['model'],
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text}
        ],
        options=options
    )
    summary_text = response['message']['content']

    # The model occasionally returns an empty answer; one retry with a higher
    # temperature usually breaks that degenerate state
    if not summary_text.strip():
        logger.warning("LLM returned an empty summary, retrying once with higher temperature")
        retry_options = dict(options, temperature=min(1.0, configs['llm']['temperature'] + 0.5))
        response = chat(
            model=configs['llm']['model'],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            options=retry_options
        )
        summary_text = response['message']['content']

    # The model tends to undershoot the requested length; ask it once to expand
    # its own draft when the result stays below the minimum
    current_words = word_count(summary_text)
    if summary_text.strip() and current_words < min_words:
        logger.debug(f"Summary too short ({current_words} < {min_words} words), asking the model to expand it")
        response = chat(
            model=configs['llm']['model'],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text},
                {"role": "assistant", "content": summary_text},
                {"role": "user", "content": (
                    f"Your summary has only {current_words} words. Expand it to at least "
                    f"{min_words} words by adding more detail from the transcript. Keep the "
                    f"same style and rules. Output only the expanded summary."
                )}
            ],
            options=options
        )
        expanded_text = response['message']['content']
        if word_count(expanded_text) > current_words:
            summary_text = expanded_text

    return summary_text


def hierarchical_reduce(texts, prompt_hint, summary_language, min_words, max_words, prompt_template, max_workers=4):
    while len(texts) > 1:
        groups = group_texts_by_tokens(texts, llm_input_token_budget)
        final_round = len(groups) == 1

        new_texts = [None] * len(groups)
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {}
            for idx, group in enumerate(groups):
                group_text = "\n".join(group)
                if final_round:
                    group_min_words, group_max_words = min_words, max_words
                else:
                    group_target_words = min(800, max(250, int(0.25 * word_count(group_text))))
                    group_min_words = int(group_target_words * 0.8)
                    group_max_words = int(group_target_words * 1.2)
                futures[executor.submit(
                    summarize_chunk, group_text, prompt_hint, summary_language,
                    group_min_words, group_max_words, prompt_template
                )] = idx
            for future in futures:
                new_texts[futures[future]] = future.result()

        texts = new_texts

    return texts[0]


def summarize_file(lines, plain_text, obfuscated_stem, prompt_hint, summary_language,
                   prompt_template, docx_title, filename_postfix):
    total_words = word_count(plain_text)
    target_words = min(1500, max(150, int(0.2 * total_words)))
    min_words = int(target_words * 0.8)
    max_words = int(target_words * 1.2)

    chunks = chunk_lines_by_tokens(lines, llm_input_token_budget)
    logger.debug(f"Chunks: {len(chunks)}, target summary length: {min_words}-{max_words} words")

    if len(chunks) == 1:
        final_text = summarize_chunk(chunks[0], prompt_hint, summary_language,
                                     min_words, max_words, prompt_template)
    else:
        final_text = hierarchical_reduce(chunks, prompt_hint, summary_language,
                                         min_words, max_words, prompt_template) 

    final_text = clean_llm_output(final_text)

    summary_path = os.path.join(dir_processed, obfuscated_stem + filename_postfix + ".docx")
    document = Document()
    document.add_heading(docx_title, 0)
    document.add_paragraph(final_text)
    document.save(summary_path)
    return summary_path



def transcribe_file(obfuscated_standardized_fullpath, sidecar_path):
    logger.debug("Transcribing file...")
    obfuscated_standardized_fullpath = pathlib.Path(obfuscated_standardized_fullpath)
    obfuscated_diarization_wav_fullpath = None

    try:
        # Remember start time of the transcription process
        transcription_start_time = time.time()

        # Clarify obfuscated filename stem
        obfuscated_stem = obfuscated_standardized_fullpath.stem
        structured_filename = mws_helpers.clarify_string(obfuscated_stem)

        # Extract Language Code
        language_code = mws_helpers.get_language_setting_index_or_code(int(structured_filename.split('#', 9)[3]))
        # Extract Translation Status
        translation_status = int(structured_filename.split('#', 9)[4])
        translation_status = 'translate' if translation_status == 1 else None
        # Extract Diarization Setting
        diarization_setting = int(structured_filename.split('#', 9)[5])
        # Extract Selected Transcription Model
        selected_transcription_model = mws_helpers.get_model_setting_index_or_name(
            int(structured_filename.split('#', 9)[8]))

        # Read Opus duration and size
        file_duration = OggOpus(str(obfuscated_standardized_fullpath)).info.length
        file_size = os.path.getsize(obfuscated_standardized_fullpath)

        # Extract subtitle Setting from Base Name
        subtitle_setting = int(os.path.basename(structured_filename).split('#', 9)[6])

        summary_setting = int(os.path.basename(structured_filename).split('#', 9)[7])
        # Load the model
        logger.debug("Loading Whisper...")
        if torch.cuda.is_available():
            whisper_model = whisper.load_model(selected_transcription_model).cuda().eval()
        else:
            whisper_model = whisper.load_model(selected_transcription_model)

        # Transcribe
        logger.debug(f"Transcription starts for {obfuscated_standardized_fullpath}")

        result = whisper_model.transcribe(
            str(obfuscated_standardized_fullpath),
            verbose=True,
            word_timestamps=True,
            language=language_code,
            task=translation_status
        )

        # Create subtitles if requested
        subtitle_srt_file = None
        subtitle_vtt_file = None

        if subtitle_setting == 1:
            # Write subtitles under the obfuscated stem, like the Word documents, so that the
            # clarified name (which contains the uploader's email address) never lands on disk
            srt_writer = get_writer("srt", dir_processed)
            srt_writer(result, str(obfuscated_standardized_fullpath))

            vtt_writer = get_writer("vtt", dir_processed)
            vtt_writer(result, str(obfuscated_standardized_fullpath))

            subtitle_srt_file = os.path.join(
                dir_processed,
                obfuscated_stem + ".srt"
            )

            subtitle_vtt_file = os.path.join(
                dir_processed,
                obfuscated_stem + ".vtt"
            )

        # Prepare full name and create document
        logger.debug("Creating Word Document...")

        transcript_text_only_file_fullname = os.path.join(
            dir_processed,
            obfuscated_stem
            + configs['texts']['whisper']['text_only_attachment_postfix']
            + '.docx'
        )

        document_text_only = Document()

        # Change Docx Settings
        style = document_text_only.styles['Normal']
        font = style.font
        font.name = configs['ui_settings']['corporate_design_font']
        font.size = Pt(10)

        # Define confidence levels bounds
        high_confidence_level_lower_bound = configs['features']['high_confidence_level_lower_bound']
        average_confidence_level_lower_bound = configs['features']['average_confidence_level_lower_bound']

        # Add colors legend
        document_text_only.add_heading(configs['texts']['whisper']['docx_color_legends_label'])

        confidence_high_run = document_text_only.add_paragraph().add_run(
            f"Hohes Konfidenzniveau (> {high_confidence_level_lower_bound}): "
            f"Text wird nicht farblich hervorgehoben."
        )

        confidence_average_run = document_text_only.add_paragraph().add_run(
            f"Mittleres Konfidenzniveau "
            f"({average_confidence_level_lower_bound} ≤ Wert ≤ {high_confidence_level_lower_bound}): "
            f"Text wird gelb hervorgehoben."
        )
        confidence_average_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        confidence_low_run = document_text_only.add_paragraph().add_run(
            f"Niedriges Konfidenzniveau (< {average_confidence_level_lower_bound}): "
            f"Text wird rot hervorgehoben."
        )
        confidence_low_run.font.highlight_color = WD_COLOR_INDEX.RED

        # Add heading
        document_text_only.add_heading(configs['texts']['whisper']['docx_text_only_title'])

        # Add new paragraph
        text_paragraph = document_text_only.add_paragraph()

        # Function for coloring
        def select_highlight_color(confidence_value):
            if confidence_value > high_confidence_level_lower_bound:
                return None
            elif average_confidence_level_lower_bound <= confidence_value <= high_confidence_level_lower_bound:
                return WD_COLOR_INDEX.YELLOW
            else:
                return WD_COLOR_INDEX.RED

        # Add Text Segments in Loop
        if "segments" in result:
            for segment in result["segments"]:
                if "words" in segment:
                    segment_confidency = segment['avg_logprob']
                    print(segment_confidency)

                    highlight_color = select_highlight_color(segment_confidency)
                    run_text = segment['text']

                    current_run = text_paragraph.add_run(run_text)
                    current_run.font.highlight_color = highlight_color

        # Save file
        document_text_only.save(transcript_text_only_file_fullname)

        # Perform diarization conditionally
        completed_conversation_turns = None
        if diarization_setting == 1:
            # Pyannote/torchaudio can be more reliable with WAV than Opus.
            # So we create a temporary WAV only for diarization.
            obfuscated_diarization_wav_fullpath = pathlib.Path(
                dir_in_progress,
                obfuscated_stem + "_diarization.wav"
            )

            diarization_stream = ffmpeg.input(str(obfuscated_standardized_fullpath)).audio

            diarization_stream = ffmpeg.output(
                diarization_stream,
                str(obfuscated_diarization_wav_fullpath),
                **{
                    "ac": 1,
                    "ar": 16000,
                    "c:a": "pcm_s16le",
                }
            )

            try:
                ffmpeg.run(
                    diarization_stream,
                    overwrite_output=True,
                    capture_stdout=True,
                    capture_stderr=True
                )
            except ffmpeg.Error as ffmpeg_error:
                stderr = (
                    ffmpeg_error.stderr.decode("utf-8", errors="replace")
                    if ffmpeg_error.stderr
                    else ""
                )
                logger.error(f"FFmpeg WAV conversion for diarization failed:\n{stderr}", exc_info=True)
                raise RuntimeError(f"FFmpeg WAV conversion for diarization failed:\n{stderr}") from ffmpeg_error

            conversation_turns_diarized = diarize_file(str(obfuscated_diarization_wav_fullpath))

            # Delete temporary WAV after successful diarization
            mws_helpers.safe_unlink(obfuscated_diarization_wav_fullpath, "temporary diarization WAV file")
            obfuscated_diarization_wav_fullpath = None

            # Normalize diarized turns
            normalized_and_diarized_turns = []
            processed_turns_indexes = []

            start_clipboarded = None
            end_clipboarded = None
            speaker_clipboarded = None

            for main_turn_checked in conversation_turns_diarized:
                # Check if this turn has already been normalized
                if not main_turn_checked['index'] in processed_turns_indexes:
                    # If not, remember its starting time
                    start_clipboarded = main_turn_checked['start']

                    # Remember its ending time
                    end_clipboarded = main_turn_checked['end']

                    # Remember its speaker
                    speaker_clipboarded = main_turn_checked['speaker']

                    # Extract only unprocessed turns for counter-check
                    counterchecked_turns = []

                    for counterchecked_turn in conversation_turns_diarized:
                        if counterchecked_turn['index'] > main_turn_checked['index']:
                            counterchecked_turns.append(counterchecked_turn)

                    # Perform counter-check
                    for counterchecked_turn in counterchecked_turns:
                        if counterchecked_turn['speaker'] == main_turn_checked['speaker']:
                            end_clipboarded = counterchecked_turn['end']
                            processed_turns_indexes.append(counterchecked_turn['index'])
                        else:
                            processed_turns_indexes.append(main_turn_checked['index'])
                            break

                    # Protocol normalized turn
                    normalized_and_diarized_turns.append({
                        'start': start_clipboarded,
                        'end': end_clipboarded,
                        'speaker': speaker_clipboarded
                    })

            # Prepare word-level timestamps
            timestamped_words = []

            for segment in result['segments']:
                for word in segment['words']:
                    timestamped_words.append({
                        'word': word['word'],
                        'start': word['start'],
                        'end': word['end']
                    })

            # Assign words to conversation turns
            completed_conversation_turns = diarize_timestamped_words(
                normalized_and_diarized_turns,
                timestamped_words
            )

            # Save conversation turns to a Docx File
            def convert_secs_to_timestamp(seconds):
                seconds_in_hours = seconds / 3600
                timestamp = str(datetime.timedelta(hours=seconds_in_hours))[:10]

                if len(timestamp) == 7:
                    timestamp = timestamp + '.00'

                return timestamp

            transcript_conversation_turns_file_fullname = os.path.join(
                dir_processed,
                obfuscated_stem
                + configs['texts']['whisper']['conversation_turns_attachment_postfix']
                + '.docx'
            )

            # Create Docx file
            document_conversation_turns = Document()

            # Change Docx Settings
            style = document_conversation_turns.styles['Normal']
            font = style.font
            font.name = configs['ui_settings']['corporate_design_font']
            font.size = Pt(10)

            # Add heading
            document_conversation_turns.add_heading(
                configs['texts']['whisper']['docx_text_conversation_turns_title']
            )

            # Add all conversation turns
            for conversation_turn in completed_conversation_turns:
                # Add empty speaker line
                p = document_conversation_turns.add_paragraph('')

                # Add bold speaker/timestamp line
                p.add_run(
                    f"{conversation_turn['speaker']} -> "
                    f"{convert_secs_to_timestamp(conversation_turn['start'])}-"
                    f"{convert_secs_to_timestamp(conversation_turn['end'])}"
                ).bold = True

                # Add conversation text
                document_conversation_turns.add_paragraph(
                    f"{conversation_turn['text']}"
                )

                # Add empty line
                document_conversation_turns.add_paragraph('')

            document_conversation_turns.save(transcript_conversation_turns_file_fullname)

        else:
            transcript_conversation_turns_file_fullname = None

        # Create summaries if requested
        summary_file = None
        conversation_summary_file = None
        if summary_setting == 1:
            # Fall back to defaults if the sidecar is missing, so a finished transcription
            # is never lost just because the optional summary options could not be read
            sidecar = {}
            if sidecar_path.exists():
                with open(sidecar_path, 'r', encoding='utf-8') as f:
                    sidecar = json.load(f)
            else:
                logger.warning(f"Sidecar file not found, summarizing with default options: {sidecar_path}")

            prompt_hint = sidecar.get("prompt_hint") or ""
            summary_language = (sidecar.get("summary_language") or
                                configs['texts']['page']['summary_language_code_selectbox_default_option'])

            logger.debug("Creating Summary...")
            summary_file = summarize_file(
                format_segments_for_llm(result["segments"]),
                result["text"],
                obfuscated_stem,
                prompt_hint,
                summary_language,
                configs['texts']['whisper']['summary_prompt'],
                configs['texts']['whisper']['docx_summary_title'],
                configs['texts']['whisper']['summary_attachment_postfix']
            )

            if completed_conversation_turns is not None:
                logger.debug("Creating Conversation Summary...")
                conversation_summary_file = summarize_file(
                    format_conversation_turns_for_llm(completed_conversation_turns),
                    result["text"],
                    obfuscated_stem,
                    prompt_hint,
                    summary_language,
                    configs['texts']['whisper']['conversation_summary_prompt'],
                    configs['texts']['whisper']['docx_conversation_summary_title'],
                    configs['texts']['whisper']['conversation_summary_attachment_postfix']
                )

        # Delete the standardized Opus file after transcription is complete
        mws_helpers.safe_unlink(obfuscated_standardized_fullpath, "standardized Opus file")

        # Get ending time of the transcription
        transcription_end_time = time.time()

        # Prepare new line for the performance stats
        language_code_for_protocol = '--' if language_code is None else language_code
        translation_status_for_protocol = 'translate' if translation_status == "translate" else 'original'

        new_perf_record = [{
            'model': selected_transcription_model,
            'language_code': language_code_for_protocol,
            'translation_status': translation_status_for_protocol,
            'diarization_status': diarization_setting,
            'duration_seconds': file_duration,
            'file_size': file_size,
            'transcription_start_time': transcription_start_time,
            'transcription_end_time': transcription_end_time,
            'transcription_time_per_one_raw_second': (
                transcription_end_time - transcription_start_time
            ) / file_duration
        }]

        # Create a dataframe of this dictionary
        new_perf_record_df = pandas.DataFrame(new_perf_record)

        # Load previous performance protocol
        perf_protocol = pandas.read_csv(path_to_perf_protocol, encoding='Windows-1252')

        # Concatenate both frames
        perf_protocol_updated = pandas.concat(
            [perf_protocol, new_perf_record_df],
            ignore_index=True
        )

        # Save new state of the protocol
        perf_protocol_updated.to_csv(
            path_to_perf_protocol,
            encoding='Windows-1252',
            index=False
        )

        return [
            transcript_text_only_file_fullname,
            transcript_conversation_turns_file_fullname,
            file_duration,
            file_size,
            subtitle_vtt_file,
            subtitle_srt_file,
            summary_file,
            conversation_summary_file
        ]

    except Exception:
        logger.critical('Transcription failed.', exc_info=True)
        # Get exception infos
        error_string = traceback.format_exc()
        error_message_for_admins = f"({getpass.getuser()}) - {error_string}"
        notify_admins(error_message_for_admins)

        # Cleanup temporary files
        mws_helpers.safe_unlink(obfuscated_diarization_wav_fullpath, "temporary diarization WAV file")
        mws_helpers.safe_unlink(obfuscated_standardized_fullpath, "standardized Opus file")

        notify_admins(
            f"({getpass.getuser()}) File that resulted in error has been deleted"
        )

        # Re-raise so process_file also knows that transcription failed
        raise


def get_decrypted_bytes(enc_file_fullpath):
    from cryptography.fernet import Fernet

    # Get key
    key = mws_helpers.get_encryption_key()
    fernet = Fernet(key)

    # Read the encrypted bytes from the file
    with open(enc_file_fullpath, "rb") as enc_file:
        encrypted_bytes = enc_file.read()

    # Decrypt the bytes
    decrypted_bytes = fernet.decrypt(encrypted_bytes)

    return decrypted_bytes

def process_file(obfuscated_encrypted_fullpath, processing_marker_fullpath=None):
    file_path = Path(obfuscated_encrypted_fullpath)
    sidecar_path = file_path.with_suffix('.json')
    logger.debug('Processing file...')
    obfuscated_decrypted_fullpath = None
    obfuscated_standardized_fullpath = None
    obfuscated_filename_stem = None

    transcript_text_only_file_fullname = None
    transcript_conversation_turns_file_fullname = None
    subtitle_vtt_file_fullname = None
    subtitle_srt_file_fullname = None
    summary_file_fullname = None
    conversation_summary_file_fullname = None

    try:
        # Prepare fullpath for conversion folder
        new_location_fullpath = os.path.join(
            dir_format_conversion,
            pathlib.Path(obfuscated_encrypted_fullpath).name
        )

        # Move original file to the conversion folder
        os.replace(obfuscated_encrypted_fullpath, new_location_fullpath)
        obfuscated_encrypted_fullpath = new_location_fullpath

        obfuscated_filename_stem = pathlib.Path(obfuscated_encrypted_fullpath).stem

        # Decrypt original file bytes
        decrypted_original_file_bytes = get_decrypted_bytes(obfuscated_encrypted_fullpath)

        # Overwrite same file with decrypted data
        with open(obfuscated_encrypted_fullpath, "wb") as f:
            f.write(decrypted_original_file_bytes)

        obfuscated_decrypted_fullpath = obfuscated_encrypted_fullpath

        # Write a standardized small Opus file
        obfuscated_standardized_fullpath = pathlib.Path(
            dir_in_progress,
            obfuscated_filename_stem + ".opus"
        )

        stream = ffmpeg.input(obfuscated_decrypted_fullpath).audio

        stream = ffmpeg.output(
            stream,
            str(obfuscated_standardized_fullpath),
            **{
                "ac": 1,
                "ar": 16000,
                "c:a": "libopus",
                "b:a": "24k",
            }
        )

        try:
            ffmpeg.run(
                stream,
                overwrite_output=True,
                capture_stdout=True,
                capture_stderr=True
            )
        except ffmpeg.Error as ffmpeg_error:
            stderr = (
                ffmpeg_error.stderr.decode("utf-8", errors="replace")
                if ffmpeg_error.stderr
                else ""
            )
            logger.error(f"FFmpeg Opus conversion failed:\n{stderr}", exc_info=True)
            raise RuntimeError(f"FFmpeg Opus conversion failed:\n{stderr}") from ffmpeg_error

        # Delete decrypted original only after successful conversion
        mws_helpers.safe_unlink(obfuscated_decrypted_fullpath, "decrypted original file")
        obfuscated_decrypted_fullpath = None

        # Get starting time
        loop_start_time = time.time()

        # Start transcription
        transcription_result_paths = transcribe_file(obfuscated_standardized_fullpath, sidecar_path)
        if not transcription_result_paths:
            logger.error("transcribe_file did not return result paths", exc_info=True)
            raise RuntimeError("transcribe_file did not return result paths")
        transcript_text_only_file_fullname = transcription_result_paths[0]
        transcript_conversation_turns_file_fullname = transcription_result_paths[1]
        duration_seconds = transcription_result_paths[2]
        file_size = transcription_result_paths[3]
        subtitle_vtt_file_fullname = transcription_result_paths[4]
        subtitle_srt_file_fullname = transcription_result_paths[5]
        summary_file_fullname = transcription_result_paths[6]
        conversation_summary_file_fullname = transcription_result_paths[7]

        # transcribe_file deletes the standardized Opus file after successful transcription
        obfuscated_standardized_fullpath = None

        # Gather file info for message
        if not duration_seconds:
            duration_minutes = 0
            message_text_for_later = "Could not read duration for file"
            logger.warning(message_text_for_later)
        else:
            duration_minutes = round(duration_seconds / 60, 2)
            message_text_for_later = f"{duration_minutes} min. long"

        # Get finish time
        loop_finish_time = time.time()
        elapsed_time = loop_finish_time - loop_start_time
        time_used_to_transcribe = str(datetime.timedelta(seconds=elapsed_time))

        # Prepare message
        if transcript_conversation_turns_file_fullname is not None:
            email_text_key_parts = ['email_text_two_files']
        else:
            email_text_key_parts = ['email_text_one_file']

        if subtitle_srt_file_fullname is not None or subtitle_vtt_file_fullname is not None:
            email_text_key_parts.append('subtitles')

        if summary_file_fullname is not None:
            email_text_key_parts.append('summary')

        email_text_key = '_'.join(email_text_key_parts)

        # Fall back to the base text if the config file does not know the key yet
        if email_text_key not in configs['texts']['whisper']:
            logger.warning(f"Missing email text '{email_text_key}' in config, using fallback")

        email_text = configs['texts']['whisper'].get(
            email_text_key,
            configs['texts']['whisper'][email_text_key_parts[0]]
        )

        # Extract Email Address and File Name from Base Name
        clarified_stem = mws_helpers.clarify_string(obfuscated_filename_stem)
        email_address = clarified_stem.split('#', 8)[2]
        file_name = clarified_stem.split('#', 9)[9]

        # Prepare E-Mail subject
        subject_ready = f"{configs['texts']['whisper']['email_subject']}: {file_name}"

        if len(subject_ready) > 75:
            email_subject = subject_ready[:75] + '...'
        else:
            email_subject = subject_ready

        # Prepare attachments with custom email filenames
        attachments = []

        # Text-only transcript
        attachments.append(
            (
                transcript_text_only_file_fullname,
                f"{file_name}{configs['texts']['whisper']['text_only_attachment_postfix']}.docx"
            )
        )

        # Conversation turns transcript, optional
        if transcript_conversation_turns_file_fullname is not None:
            attachments.append(
                (
                    transcript_conversation_turns_file_fullname,
                    f"{file_name}{configs['texts']['whisper']['conversation_turns_attachment_postfix']}.docx"
                )
            )

        if summary_file_fullname is not None:
            attachments.append(
                (
                    summary_file_fullname,
                    f"{file_name}{configs['texts']['whisper']['summary_attachment_postfix']}.docx"
                )
            )

        # Conversation summary, optional
        if conversation_summary_file_fullname is not None:
            attachments.append(
                (
                    conversation_summary_file_fullname,
                    f"{file_name}{configs['texts']['whisper']['conversation_summary_attachment_postfix']}.docx"
                )
            )

        # Subtitle files, optional
        if subtitle_srt_file_fullname is not None:
            attachments.append(
                (
                    subtitle_srt_file_fullname,
                    f"{file_name}{configs['texts']['whisper']['subtitle_attachment_postfix']}.srt"
                )
            )

        if subtitle_vtt_file_fullname is not None:
            attachments.append(
                (
                    subtitle_vtt_file_fullname,
                    f"{file_name}{configs['texts']['whisper']['subtitle_attachment_postfix']}.vtt"
                )
            )

        # Send the results of transcribing
        try:
            mws_helpers.send_mail(
                configs['email']['noreply_email'],
                [email_address],
                email_subject,
                email_text,
                attachments
            )

        except Exception as e:
            logger.error('Could not send email', exc_info=True)
            notify_admins(
                f"({getpass.getuser()}) - Transcription was successful, "
                f"but an error occurred when trying to send the email"
            )

            e_type, e_object, e_traceback = sys.exc_info()
            e_line_number = e_traceback.tb_lineno

            error_message_for_admins = (
                f"({getpass.getuser()}) - Following error happened when trying to send the email: "
                f"{e.__class__.__name__}."
                f"{new_line_for_f_strings}{new_line_for_f_strings}"
                f"Error raised on line: {e_line_number}"
                f"{new_line_for_f_strings}{new_line_for_f_strings}"
                f"{e}"
            )

            notify_admins(error_message_for_admins)

            # Copy transcription results to local testing folder
            files_list = [
                transcript_text_only_file_fullname,
                transcript_conversation_turns_file_fullname,
                subtitle_vtt_file_fullname,
                subtitle_srt_file_fullname,
                summary_file_fullname,
                conversation_summary_file_fullname
            ]

            files_list = [f for f in files_list if f is not None]

            for results_file in files_list:
                try:
                    shutil.copy(
                        results_file,
                        os.path.join(
                            mws_helpers.ProjectPaths().local_tests_folder_path,
                            os.path.basename(results_file)
                        )
                    )
                except FileNotFoundError:
                    logger.error('Source file not found!', exc_info=True)
                except PermissionError:
                    logger.error('Permission denied!', exc_info=True)
                except Exception as copy_error:
                    logger.error(f'An error occurred while copying {results_file}: {copy_error}', exc_info=True)

        # Delete Word and subtitle files after sending/copying them
        mws_helpers.safe_unlink(transcript_text_only_file_fullname, "text-only transcript DOCX file")
        transcript_text_only_file_fullname = None

        mws_helpers.safe_unlink(
            transcript_conversation_turns_file_fullname,
            "conversation-turns transcript DOCX file"
        )
        transcript_conversation_turns_file_fullname = None

        mws_helpers.safe_unlink(subtitle_srt_file_fullname, "SRT subtitle file")
        subtitle_srt_file_fullname = None

        mws_helpers.safe_unlink(subtitle_vtt_file_fullname, "VTT subtitle file")
        subtitle_vtt_file_fullname = None

        mws_helpers.safe_unlink(summary_file_fullname, "SUMMARY file")
        summary_file_fullname = None

        mws_helpers.safe_unlink(conversation_summary_file_fullname, "CONVERSATION SUMMARY file")
        conversation_summary_file_fullname = None

        mws_helpers.safe_unlink(sidecar_path, "Sidecar file")
        sidecar_path = None

        # Send notification
        count_unprocessed, _ = mws_helpers.count_and_list_files(dir_orig_files_temps)
        count_in_progress, _ = mws_helpers.count_processing_jobs()
        notify_admins(
            f'({getpass.getuser()}) - A file has been successfully transcribed '
            f'({message_text_for_later})\n'
            f"Files Waiting: {count_unprocessed}\n"
            f"Files in Progress: {count_in_progress}/{configs['features']['max_files_processed_simultaneously']}\n"

        )

    except Exception:
        logger.critical('Could not process file!', exc_info=True)
        error_string = traceback.format_exc()
        error_message_for_admins = f"({getpass.getuser()}) - {error_string}"
        notify_admins(error_message_for_admins)

        # Inform the user by email that processing failed. This covers all failure
        # modes (decryption, FFmpeg conversion, transcription)
        try:
            error_mail_clarified_stem = mws_helpers.clarify_string(obfuscated_filename_stem)
            error_mail_address = error_mail_clarified_stem.split('#', 8)[2]

            error_email_subject = configs['texts']['whisper']['email_subject_error']
            error_email_text = configs['texts']['whisper']['email_text_error']

            mws_helpers.send_mail(
                configs['email']['noreply_email'],
                [error_mail_address],
                error_email_subject,
                error_email_text
            )

        except Exception as error_mail_exception:
            logger.error(f"Could not send error notification email: {error_mail_exception}", exc_info=True)

        # Cleanup files that may have remained after an error
        mws_helpers.safe_unlink(obfuscated_decrypted_fullpath, "decrypted original file")
        mws_helpers.safe_unlink(obfuscated_standardized_fullpath, "standardized Opus file")

        mws_helpers.safe_unlink(transcript_text_only_file_fullname, "text-only transcript DOCX file")
        mws_helpers.safe_unlink(
            transcript_conversation_turns_file_fullname,
            "conversation-turns transcript DOCX file"
        )
        mws_helpers.safe_unlink(subtitle_srt_file_fullname, "SRT subtitle file")
        mws_helpers.safe_unlink(subtitle_vtt_file_fullname, "VTT subtitle file")

        mws_helpers.safe_unlink(summary_file_fullname, "SUMMARY file")
        mws_helpers.safe_unlink(conversation_summary_file_fullname, "CONVERSATION SUMMARY file")
        mws_helpers.safe_unlink(sidecar_path, "Sidecar file")

    finally:
        # This is the important part:
        # one finished daemon process = remove one active-job marker.
        mws_helpers.safe_unlink(processing_marker_fullpath)


def main():
    from multiprocessing import Process

    # Clean old job markers once when the daemon starts.
    # This prevents old .job files from blocking processing after a restart.
    mws_helpers.cleanup_processing_markers()

    # Infinite Loop
    while 1 < 2:
        seconds = 10
        # Count active jobs by .job marker files only
        count_files_in_proggress, _ = mws_helpers.count_processing_jobs()
        if count_files_in_proggress < configs['features']['max_files_processed_simultaneously']:
            # List unprocessed files
            count_unprocessed, unprocessed_files = mws_helpers.count_and_list_files(
                dir_temp_orig_files
            )

            if count_unprocessed >= 1:
                fullname_of_next_unprocessed_file = unprocessed_files[0]

                # Create the processing marker BEFORE starting the worker process.
                # This makes the count increase immediately.
                processing_marker_fullpath = mws_helpers.create_processing_marker(
                    fullname_of_next_unprocessed_file
                )

                another_daemon_process = Process(
                    target=process_file,
                    args=(
                        fullname_of_next_unprocessed_file,
                        processing_marker_fullpath
                    )
                )

                another_daemon_process.daemon = True

                try:
                    another_daemon_process.start()
                except Exception:
                    # If the process could not start, remove the marker again.
                    mws_helpers.safe_unlink(processing_marker_fullpath)
                    logger.error('Could not start daemon process', exc_info=True)
                    raise

                logger.debug(
                    f"Something has been loaded and we created a new daemon process for it! "
                    f"Active jobs: {count_files_in_proggress + 1}. "
                    f"Let's sleep again for {seconds} seconds till the next check..."
                )

            else:
                logger.debug(
                    f"Well... Nothing was loaded in the meanwhile! "
                    f"Active jobs: {count_files_in_proggress}. "
                    f"Let's sleep again for {seconds} seconds till the next check..."
                )

        time.sleep(seconds)


if __name__ == "__main__":
    main()